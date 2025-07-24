import streamlit as st
from openai import Client
from docx import Document
import pdfplumber
import os
from fpdf import FPDF
import base64

from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

client = Client()
client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY")    
)

# Set OpenAI API Key
#openai.api_key = ""  # Replace with your OpenAI API key


# Function to extract text from DOCX
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

# Function to extract text from PDF
def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text() + '\n'
    return text

# Function for GPT-based translation
def gpt_translate(content, target_language):
    prompt = f"Translate the following text into {target_language}:\n\n{content}"
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert translator."},
            {"role": "user", "content": prompt},
        ],
        temperature=0,
    )
    return response.choices[0].message.content #response['choices'][0]['message']['content']
# Function to save text as a Word document
def save_to_word(content, filename="translated_text.docx"):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)

# Function to generate a download link for a file
def download_file(file_path, file_label):
    with open(file_path, "rb") as file:
        b64 = base64.b64encode(file.read()).decode()
        href = f'<a href="data:file/{file_path.split(".")[-1]};base64,{b64}" download="{file_path}">{file_label}</a>'
        st.markdown(href, unsafe_allow_html=True)

# Streamlit app interface
st.title("Document Translation with GPT")
st.write("Upload your document, choose a target language, and get the translation!")


# File uploader
text_area = st.text_area("Enter the text to translate")
uploaded_file = st.file_uploader("Upload a file (TXT, DOCX, PDF)", type=["txt", "docx", "pdf"])
target_language = st.text_input("Enter target language (e.g., French, English, German):")

if uploaded_file and target_language:
    if st.button("Translate"):
        # Extract text based on file type
        if uploaded_file.name.endswith(".txt"):
            content = uploaded_file.read().decode("utf-8")
        elif uploaded_file.name.endswith(".docx"):
            content = extract_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith(".pdf"):
            content = extract_text_from_pdf(uploaded_file)
        
        # Translate with GPT
        with st.spinner("Translating with GPT..."):
            translation = gpt_translate(content, target_language)
        # Display and download translation
        st.subheader("Translated Text:")
        st.text_area("Translation", translation, height=300)
            
        # Save the translated text as Word and PDF files
        save_to_word(translation, "translated_text.docx")       
       
        # Provide download links
        st.download_button(
            download_file("translated_text.docx", "Download as Word")           
        )
if text_area and target_language:
    if st.button("Translate"):
        # Translate with GPT
        with st.spinner("Translating with GPT..."):
            translation = gpt_translate(text_area, target_language)
        # Display and download translation
        st.subheader("Translated Text:")
        st.text_area("Translation", translation, height=300)
    
       
    
        
  
        
    
